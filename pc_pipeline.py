#!/usr/bin/env python3
import os
import argparse
import logging
import re
import time
import unicodedata
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import List, Dict, Generator, Optional
from dataclasses import dataclass
from hashlib import sha256
from dotenv import load_dotenv

import openai
import tiktoken
from pinecone import Pinecone, ServerlessSpec
from tqdm import tqdm

# Optional imports with fallbacks
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
try:
    from docx import Document
except ImportError:
    Document = None

# -------------------- CONFIGURATION --------------------
class Config:
    DEFAULT_CHUNK_SIZE = 512
    MAX_CHUNK_TOKENS = 8191
    CHUNK_OVERLAP = 50
    MAX_BATCH_SIZE = 100
    EMBEDDING_MODEL = "text-embedding-3-large"
    EMBEDDING_DIM = 3072
    MAX_RETRIES = 3
    RETRY_DELAY = 1.0
    TEXT_NORMALIZATION = {
        "replace_whitespace": True,
        "strip_control_chars": True,
        "normalize_unicode": True
    }

# -------------------- DATA STRUCTURES --------------------
@dataclass
class DocumentChunk:
    text: str
    metadata: Dict[str, str]
    token_count: int
    embedding: Optional[List[float]] = None

@dataclass
class ProcessedDocument:
    path: Path
    chunks: List[DocumentChunk]
    processing_time: float

# -------------------- TEXT PROCESSING --------------------
class TextProcessor:
    def __init__(self):
        self.encoder = tiktoken.get_encoding("cl100k_base")
        
    def normalize_text(self, text: str) -> str:
        if Config.TEXT_NORMALIZATION["normalize_unicode"]:
            text = unicodedata.normalize("NFKC", text)
        if Config.TEXT_NORMALIZATION["replace_whitespace"]:
            text = re.sub(r'\s+', ' ', text)
        if Config.TEXT_NORMALIZATION["strip_control_chars"]:
            text = text.encode("utf-8", "ignore").decode("utf-8")
        return text.strip()

    def semantic_chunking(self, text: str) -> Generator[str, None, None]:
        paragraphs = re.split(r'\n\s*\n', text)
        current_chunk = []
        current_count = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
                
            tokens = self.encoder.encode(para)
            token_count = len(tokens)
            
            if token_count > Config.MAX_CHUNK_TOKENS:
                logging.warning(f"Oversized paragraph: {token_count} tokens")
                continue
                
            if current_count + token_count > Config.DEFAULT_CHUNK_SIZE:
                yield " ".join(current_chunk)
                current_chunk = current_chunk[-Config.CHUNK_OVERLAP:]  
                current_count = sum(len(self.encoder.encode(s)) for s in current_chunk)
            
            current_chunk.append(para)
            current_count += token_count

        if current_chunk:
            yield " ".join(current_chunk)

# -------------------- FILE PROCESSORS --------------------
class FileProcessor:
    def __init__(self, text_processor: TextProcessor):
        self.text_processor = text_processor
        self.logger = logging.getLogger(__name__)

    def process(self, file_path: Path) -> Optional[ProcessedDocument]:
        try:
            start_time = time.time()
            
            if file_path.suffix == ".pdf":
                chunks = self._process_pdf(file_path)
            elif file_path.suffix == ".docx":
                chunks = self._process_docx(file_path)
            elif file_path.suffix in (".txt", ".md"):
                chunks = self._process_text_file(file_path)
            else:
                self.logger.warning(f"Skipping unsupported file: {file_path.name}")
                return None

            return ProcessedDocument(
                path=file_path,
                chunks=chunks,
                processing_time=time.time() - start_time
            )
            
        except Exception as e:
            self.logger.error(f"Failed to process {file_path}: {str(e)}")
            return None

    def _process_pdf(self, file_path: Path) -> List[DocumentChunk]:
        chunks = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=1, y_tolerance=1)
                if text:
                    text = self.text_processor.normalize_text(text)
                    for chunk in self.text_processor.semantic_chunking(text):
                        tokens = self.text_processor.encoder.encode(chunk)
                        chunks.append(DocumentChunk(
                            text=chunk,
                            metadata={
                                "source": str(file_path),
                                "page": str(page.page_number),
                                "file_type": "pdf"
                            },
                            token_count=len(tokens)
                        ))
        return chunks

    def _process_docx(self, file_path: Path) -> List[DocumentChunk]:
        chunks = []
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        text = self.text_processor.normalize_text(text)
        
        for chunk in self.text_processor.semantic_chunking(text):
            tokens = self.text_processor.encoder.encode(chunk)
            chunks.append(DocumentChunk(
                text=chunk,
                metadata={
                    "source": str(file_path),
                    "file_type": "docx"
                },
                token_count=len(tokens)
            ))
        return chunks

    def _process_text_file(self, file_path: Path) -> List[DocumentChunk]:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        
        text = self.text_processor.normalize_text(text)
        return [
            DocumentChunk(
                text=chunk,
                metadata={
                    "source": str(file_path),
                    "file_type": file_path.suffix[1:]
                },
                token_count=len(self.text_processor.encoder.encode(chunk))
            )
            for chunk in self.text_processor.semantic_chunking(text)
        ]

# -------------------- EMBEDDING MANAGER --------------------
class EmbeddingManager:
    def __init__(self, api_key: str):
        self.client = openai.OpenAI(api_key=api_key)
        self.logger = logging.getLogger(__name__)
        
    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        valid_texts = []
        for text in texts:
            response = self.client.embeddings.create(
                input=[text],
                model=Config.EMBEDDING_MODEL
            )
            embedding = response.data[0].embedding
            if len(embedding) != Config.EMBEDDING_DIM:
                raise ValueError(f"Invalid embedding dimension: {len(embedding)}")
            valid_texts.append(embedding)
        return valid_texts

# -------------------- PINECONE MANAGER --------------------
class PineconeManager:
    def __init__(self, api_key: str):
        self.pc = Pinecone(api_key=api_key)
        self._index_cache = {}
        self.logger = logging.getLogger(__name__)
        
    def get_index(self, index_name: str) -> Pinecone.Index:
        if index_name not in self._index_cache:
            if index_name not in self.pc.list_indexes().names():
                self.logger.info(f"Creating index: {index_name}")
                self.pc.create_index(
                    name=index_name,
                    dimension=Config.EMBEDDING_DIM,
                    metric="cosine",
                    spec=ServerlessSpec(
                        cloud="aws",
                        region="us-east-1"
                    )
                )
            # Non-gRPC connection pool configuration
            self._index_cache[index_name] = self.pc.Index(
                name=index_name,
                pool_threads=50,
                connection_pool_maxsize=50
            )
        return self._index_cache[index_name]

    def fetch_namespaces(self, index_name: str) -> List[str]:
        index = self.get_index(index_name)
        stats = index.describe_index_stats()
        return list(stats.get('namespaces', {}).keys())

    def upsert_chunks(self, index_name: str, chunks: List[DocumentChunk], namespace: Optional[str] = None):
        index = self.get_index(index_name)
        vectors = []
        
        with ThreadPoolExecutor() as executor:
            futures = []
            batch = []
            
            for chunk in chunks:
                chunk_id = sha256(chunk.text.encode()).hexdigest()
                batch.append({
                    "id": chunk_id,
                    "values": chunk.embedding,
                    "metadata": chunk.metadata
                })
                
                if len(batch) >= Config.MAX_BATCH_SIZE:
                    futures.append(executor.submit(
                        self._safe_upsert,
                        index=index,
                        batch=batch.copy(),
                        namespace=namespace
                    ))
                    batch.clear()
                    
            if batch:
                futures.append(executor.submit(
                    self._safe_upsert,
                    index=index,
                    batch=batch,
                    namespace=namespace
                ))

            for future in tqdm(futures, desc="Uploading chunks"):
                future.result()

    def _safe_upsert(self, index: Pinecone.Index, batch: List[dict], namespace: Optional[str]):
        for attempt in range(Config.MAX_RETRIES):
            try:
                return index.upsert(
                    vectors=batch,
                    namespace=namespace
                )
            except Exception as e:
                if attempt < Config.MAX_RETRIES - 1:
                    delay = Config.RETRY_DELAY * (2 ** attempt)
                    self.logger.warning(f"Upsert failed: {str(e)}. Retrying in {delay}s...")
                    time.sleep(delay)
                else:
                    self.logger.error(f"Failed upsert after {Config.MAX_RETRIES} attempts")
                    raise

# -------------------- INTERACTIVE HELPERS --------------------
def show_environment_summary(pc: Pinecone):
    print("\n" + "="*40)
    print("Current Working Directory Contents:")
    cwd = Path.cwd()
    for f in cwd.glob("*"):
        print(f"• {'DIR' if f.is_dir() else 'FILE':<4} {f.name}")
    print("\nExisting Pinecone Indexes:")
    indexes = pc.list_indexes().names() or ["No indexes found"]
    print("\n".join(f" - {idx}" for idx in indexes))
    print("="*40 + "\n")

def select_or_create_index(pc: Pinecone, provided_index: Optional[str], non_interactive: bool) -> str:
    existing_indexes = pc.list_indexes().names()
    if provided_index:
        if provided_index in existing_indexes:
            print(f"Using existing index: {provided_index}")
            return provided_index
        else:
            print(f"Creating new index: {provided_index}")
            pc.create_index(
                name=provided_index,
                dimension=Config.EMBEDDING_DIM,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1"),
            )
            return provided_index

    if non_interactive:
        default_index = "default_index"
        if default_index not in existing_indexes:
            pc.create_index(
                name=default_index,
                dimension=Config.EMBEDDING_DIM,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1"),
            )
        return default_index

    print(f"Available indexes: {', '.join(existing_indexes)}")
    while True:
        choice = input("[S]elect existing or [C]reate new? [s/c]: ").lower().strip()
        if choice == "s":
            index_name = input("Enter index name: ").strip()
            if index_name in existing_indexes:
                mgr = PineconeManager(os.getenv("PINECONE_API_KEY"))
                namespaces = mgr.fetch_namespaces(index_name)
                print(f"\nIndex '{index_name}' contains {len(namespaces)} namespaces:")
                print("\n".join(f" - {ns}" for ns in namespaces))
                return index_name
            else:
                print(f"Index '{index_name}' not found")
        elif choice == "c":
            index_name = input("New index name: ").strip()
            pc.create_index(
                name=index_name,
                dimension=Config.EMBEDDING_DIM,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1"),
            )
            return index_name

def select_namespace(provided_namespace: Optional[str], non_interactive: bool) -> Optional[str]:
    if provided_namespace:
        print(f"Using namespace: '{provided_namespace}'")
        return provided_namespace
    if non_interactive:
        return "default_namespace"
    namespace = input("Enter namespace (blank for default): ").strip()
    return namespace if namespace else None

# -------------------- MAIN PIPELINE --------------------
class DocumentPipeline:
    def __init__(self):
        self.text_processor = TextProcessor()
        self.file_processor = FileProcessor(self.text_processor)
        self.embedding_manager = EmbeddingManager(os.getenv("OPENAI_API_KEY"))
        self.pinecone_manager = PineconeManager(os.getenv("PINECONE_API_KEY"))
        
    def process_directory(self, directory: Path) -> Generator[ProcessedDocument, None, None]:
        files = [f for f in directory.rglob("*") if f.is_file()]
        
        with ThreadPoolExecutor() as executor:
            futures = [executor.submit(self.file_processor.process, f) for f in files]
            
            for future in tqdm(futures, desc="Processing files"):
                result = future.result()
                if result:
                    yield result

    def run(self, directory: Path, index_name: str, namespace: str):
        all_chunks = []
        processed_docs = self.process_directory(directory)
        
        for doc in processed_docs:
            all_chunks.extend(doc.chunks)
            self._monitor_chunk_sizes(doc.chunks)
            
        start_time = time.time()
        texts = [chunk.text for chunk in all_chunks]
        embeddings = self.embedding_manager.generate_embeddings(texts)
        self._monitor_embedding_latency(start_time)
        
        for chunk, embedding in zip(all_chunks, embeddings):
            chunk.embedding = embedding
            chunk.metadata["embedding_model"] = Config.EMBEDDING_MODEL
            
        self.pinecone_manager.upsert_chunks(
            index_name=index_name,
            chunks=all_chunks,
            namespace=namespace
        )

    def _monitor_chunk_sizes(self, chunks: List[DocumentChunk]):
        sizes = [chunk.token_count for chunk in chunks]
        avg = sum(sizes)/len(sizes)
        logging.info(f"Chunk size stats - Avg: {avg:.1f}, Min: {min(sizes)}, Max: {max(sizes)}")

    def _monitor_embedding_latency(self, start_time: float):
        duration = time.time() - start_time
        logging.info(f"Embedding generation completed in {duration:.2f} seconds")

    def query_index(self, index_name: str, query: str, top_k: int = 5, namespace: Optional[str] = None):
        """Safe query implementation with size checks"""
        if top_k > 10000:
            raise ValueError("Pinecone limits top_k to 10,000")
        
        embedding = self.embedding_manager.generate_embeddings([query])[0]
        dim = len(embedding)
        estimated_size = dim * top_k * 4  # 4 bytes per float
        
        if estimated_size > 4_000_000:
            raise ValueError(f"Estimated result size {estimated_size/1e6:.1f}MB exceeds 4MB limit")
        
        index = self.pinecone_manager.get_index(index_name)
        
        return index.query(
            vector=embedding,
            top_k=top_k,
            include_metadata=True,
            include_values=False,
            namespace=namespace,
            show_progress=True
        )

    def query_namespaces(self, index_name: str, query: str, namespaces: List[str], top_k: int = 5):
        """Multi-namespace query with safety checks"""
        if top_k > 1000:
            raise ValueError("For multi-namespace queries, top_k ≤ 1000")
            
        embedding = self.embedding_manager.generate_embeddings([query])[0]
        dim = len(embedding)
        estimated_size = dim * top_k * 4 * len(namespaces)
        
        if estimated_size > 4_000_000:
            raise ValueError(f"Estimated result size {estimated_size/1e6:.1f}MB exceeds 4MB limit")
        
        index = self.pinecone_manager.get_index(index_name)
        
        return index.query_namespaces(
            vector=embedding,
            namespaces=namespaces,
            top_k=top_k,
            include_metadata=True,
            include_values=False,
            show_progress=True
        )

# -------------------- CLI SETUP --------------------
def parse_args():
    parser = argparse.ArgumentParser(description="Document Embedding Pipeline")
    parser.add_argument("--index", type=str, help="Pinecone index name")
    parser.add_argument("--namespace", type=str, help="Pinecone namespace")
    parser.add_argument("--log-level", choices=["DEBUG", "INFO", "WARNING"], default="INFO",
                      help="Logging verbosity")
    parser.add_argument("--non-interactive", action="store_true",
                      help="Disable interactive prompts")
    return parser.parse_args()

def main():
    load_dotenv()
    args = parse_args()
    
    logging.basicConfig(
        level=args.log_level,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
    )
    
    required_vars = ["OPENAI_API_KEY", "PINECONE_API_KEY"]
    for var in required_vars:
        if not os.getenv(var):
            raise ValueError(f"{var} environment variable required")
        
    pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
    show_environment_summary(pc)
    
    index_name = select_or_create_index(pc, args.index, args.non_interactive)
    namespace = select_namespace(args.namespace, args.non_interactive)
    
    pipeline = DocumentPipeline()
    pipeline.run(Path.cwd(), index_name, namespace)
    
    if not args.non_interactive:
        while True:
            choice = input("\n[Q]uery [M]ulti-namespace [E]xit? ").lower().strip()
            if choice == "q":
                query = input("Enter search query: ")
                try:
                    results = pipeline.query_index(index_name, query, namespace=namespace)
                    print("\nTop results:")
                    for match in results.matches:
                        print(f"- {match.score:.3f}: {match.metadata['source']} (Page {match.metadata.get('page', 'N/A')})")
                except ValueError as e:
                    print(f"Query error: {str(e)}")
            elif choice == "m":
                query = input("Enter search query: ")
                namespaces = input("Enter namespaces (comma-separated): ").split(',')
                try:
                    results = pipeline.query_namespaces(index_name, query, namespaces)
                    print("\nCross-namespace results:")
                    for ns in results.results:
                        print(f"\nNamespace '{ns.namespace}':")
                        for match in ns.matches:
                            print(f"- {match.score:.3f}: {match.metadata['source']}")
                except ValueError as e:
                    print(f"Query error: {str(e)}")
            elif choice == "e":
                print("Exiting")
                break

if __name__ == "__main__":
    main()
